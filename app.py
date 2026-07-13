# -*- coding: utf-8 -*-
"""
招标文件合规检测工具 —— Flask后端
功能：文档解析 + 规则引擎检测 + LLM深度分析 + 报告生成（JSON/PDF）
"""

import os
import json
import time
import uuid
import tempfile
import traceback
from datetime import datetime

from flask import Flask, request, jsonify, render_template, send_file, Response

# 文档解析
from docx import Document
from PyPDF2 import PdfReader

# 规则引擎
from rules import run_rule_check, build_llm_prompt

# PDF报告生成
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import requests
from werkzeug.utils import secure_filename

# 设置硅基流动API Key（环境变量优先，兜底用内置Key，需在导入bid_generator前设置）
os.environ.setdefault('SILICONFLOW_API_KEY', 'sk-toiqhdxjmwxtlgzxxbbrsimrnebyfmvedsrpwesqdeyppviq')

# 投标文件生成模块
from bid_generator import (
    parse_tender_document, generate_bid_content, generate_bid_docx,
    get_material_status, get_missing_materials,
    UPLOAD_DIR, OUTPUT_DIR, BID_SESSIONS,
)
from bid_templates import (
    ALLOWED_EXTENSIONS_BID, FILE_SIZE_LIMITS, get_file_type,
    get_file_size_limit, FILE_TYPE_LABELS, MATERIAL_CATEGORIES,
)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB（支持视频上传）

# 配置
UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'docx', 'pdf', 'txt'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB（合规检测用）

# 确保目录存在
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 硅基流动API配置（Key已在文件顶部通过os.environ.setdefault设置）
SILICONFLOW_API_KEY = os.environ.get('SILICONFLOW_API_KEY', '')
SILICONFLOW_API_URL = 'https://api.siliconflow.cn/v1/chat/completions'
LLM_MODEL = 'Qwen/Qwen2.5-72B-Instruct'

# 存储检测结果（内存中，无需数据库）
REPORTS = {}

# 注册中文字体（使用系统字体）
def _register_chinese_font():
    """注册中文字体用于PDF生成"""
    font_paths = [
        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
    ]
    for path in font_paths:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont('ChineseFont', path))
                return 'ChineseFont'
            except Exception:
                continue
    # 如果没有中文字体，尝试下载文泉驿
    try:
        import subprocess
        os.system('apt-get install -y fonts-wqy-zenhei 2>/dev/null || true')
        if os.path.exists('/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc'):
            pdfmetrics.registerFont(TTFont('ChineseFont', '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc'))
            return 'ChineseFont'
    except Exception:
        pass
    return 'Helvetica'


def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ============================================================
# 文档解析模块
# ============================================================

def parse_docx(file_path):
    """解析.docx文件，提取全部文本"""
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text.strip(' |'):
                paragraphs.append(row_text)
    
    return '\n'.join(paragraphs)


def parse_pdf(file_path):
    """解析.pdf文件，提取全部文本"""
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return '\n'.join(text_parts)


def parse_txt(file_path):
    """解析.txt文件，直接读取"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def parse_document(file_path, filename):
    """
    根据文件类型选择解析器
    返回: 文档全文文本
    """
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    if ext == 'docx':
        return parse_docx(file_path)
    elif ext == 'pdf':
        return parse_pdf(file_path)
    elif ext == 'txt':
        return parse_txt(file_path)
    else:
        raise ValueError(f'不支持的文件格式: {ext}')


# ============================================================
# LLM深度合规分析
# ============================================================

def call_llm(text):
    """
    调用硅基流动API进行深度合规分析
    参数: text - 招标文件全文
    返回: 检测结果列表（可能为空）
    """
    if not SILICONFLOW_API_KEY:
        print('[INFO] 未配置SILICONFLOW_API_KEY，跳过LLM深度分析')
        return []
    
    system_prompt, user_prompt = build_llm_prompt(text)
    
    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {SILICONFLOW_API_KEY}'
    }
    
    payload = {
        'model': LLM_MODEL,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.1,  # 低温度确保结果稳定
        'max_tokens': 4096,
        'response_format': {'type': 'json_object'}
    }
    
    try:
        print('[INFO] 调用硅基流动API进行深度分析...')
        response = requests.post(SILICONFLOW_API_URL, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        
        result = response.json()
        content = result['choices'][0]['message']['content']
        
        # 解析JSON响应
        llm_result = json.loads(content)
        items = llm_result.get('items', [])
        
        # 验证并补充字段
        valid_items = []
        for item in items:
            if isinstance(item, dict) and 'description' in item:
                valid_item = {
                    'level': item.get('level', '中'),
                    'category': item.get('category', '其他风险'),
                    'description': item.get('description', ''),
                    'original_text': item.get('original_text', ''),
                    'legal_basis': item.get('legal_basis', ''),
                    'suggestion': item.get('suggestion', '')
                }
                valid_items.append(valid_item)
        
        print(f'[INFO] LLM分析完成，发现{len(valid_items)}个问题')
        return valid_items
        
    except requests.exceptions.Timeout:
        print('[WARN] LLM调用超时，跳过深度分析')
        return []
    except requests.exceptions.RequestException as e:
        print(f'[WARN] LLM调用失败: {e}')
        return []
    except (json.JSONDecodeError, KeyError, IndexError) as e:
        print(f'[WARN] LLM响应解析失败: {e}')
        return []


# ============================================================
# 报告去重与合并
# ============================================================

def merge_and_deduplicate(rule_items, llm_items):
    """
    合并规则引擎和LLM的检测结果，去重
    规则引擎的结果优先保留，LLM结果补充规则引擎未覆盖的问题
    """
    all_items = list(rule_items)
    
    # 简单去重：如果LLM结果的description与已有结果的description相似度较高，则跳过
    existing_descs = [item['description'] for item in all_items]
    
    for llm_item in llm_items:
        # 检查是否与已有结果重复（简单关键词匹配）
        is_duplicate = False
        for existing_desc in existing_descs:
            # 如果描述中有超过50%的关键词重复，则认为重复
            llm_words = set(llm_item.get('description', '').split())
            existing_words = set(existing_desc.split())
            if llm_words and existing_words:
                overlap = len(llm_words & existing_words) / max(len(llm_words), len(existing_words))
                if overlap > 0.5:
                    is_duplicate = True
                    break
        
        if not is_duplicate:
            all_items.append(llm_item)
    
    return all_items


# ============================================================
# PDF报告生成
# ============================================================

def generate_pdf_report(report_data, output_path):
    """
    生成PDF格式的检测报告
    参数: report_data - 检测报告数据(dict)
          output_path - PDF文件输出路径
    """
    font_name = _register_chinese_font()
    
    # 创建PDF文档
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=20*mm,
        leftMargin=20*mm,
        topMargin=20*mm,
        bottomMargin=20*mm
    )
    
    # 样式定义
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'ChineseTitle',
        parent=styles['Title'],
        fontName=font_name,
        fontSize=20,
        alignment=TA_CENTER,
        spaceAfter=6*mm,
        textColor=HexColor('#0a0a1a')
    )
    
    subtitle_style = ParagraphStyle(
        'ChineseSubtitle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        alignment=TA_CENTER,
        textColor=HexColor('#666666'),
        spaceAfter=10*mm
    )
    
    heading_style = ParagraphStyle(
        'ChineseHeading',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=14,
        spaceBefore=6*mm,
        spaceAfter=3*mm,
        textColor=HexColor('#0a0a1a')
    )
    
    normal_style = ParagraphStyle(
        'ChineseNormal',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=16,
        spaceAfter=2*mm
    )
    
    label_style = ParagraphStyle(
        'ChineseLabel',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=16,
        textColor=HexColor('#333333'),
        spaceAfter=1*mm
    )
    
    # 风险等级颜色
    level_colors = {
        '高': HexColor('#dc3545'),
        '中': HexColor('#fd7e14'),
        '低': HexColor('#28a745')
    }
    
    story = []
    
    # 标题
    story.append(Paragraph('NEXUS · 招标文件合规检测报告', title_style))
    story.append(Paragraph(f'生成时间：{report_data["check_time"]}', subtitle_style))
    
    # 文件信息
    story.append(Paragraph('一、文件信息', heading_style))
    info_data = [
        ['文件名称', report_data['filename']],
        ['检测时间', report_data['check_time']],
        ['检测工具', 'NEXUS招标文件合规检测系统'],
    ]
    info_table = Table(info_data, colWidths=[40*mm, 120*mm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BACKGROUND', (0, 0), (0, -1), HexColor('#f0f0f0')),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#cccccc')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 5*mm))
    
    # 检测概览
    story.append(Paragraph('二、检测概览', heading_style))
    summary = report_data['summary']
    summary_data = [
        ['总问题数', str(summary['total'])],
        ['高风险', str(summary['high'])],
        ['中风险', str(summary['medium'])],
        ['低风险', str(summary['low'])],
    ]
    summary_table = Table(summary_data, colWidths=[40*mm, 30*mm])
    summary_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BACKGROUND', (0, 0), (0, -1), HexColor('#f0f0f0')),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#cccccc')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TEXTCOLOR', (1, 1), (1, 1), level_colors['高']),
        ('TEXTCOLOR', (1, 2), (1, 2), level_colors['中']),
        ('TEXTCOLOR', (1, 3), (1, 3), level_colors['低']),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 5*mm))
    
    # 详细检测结果
    story.append(Paragraph('三、详细检测结果', heading_style))
    
    if not report_data['items']:
        story.append(Paragraph('未检测到合规问题，招标文件整体合规性良好。', normal_style))
    else:
        # 按风险等级排序：高 > 中 > 低
        level_order = {'高': 0, '中': 1, '低': 2}
        sorted_items = sorted(report_data['items'], key=lambda x: level_order.get(x['level'], 3))
        
        for i, item in enumerate(sorted_items, 1):
            level = item['level']
            level_color = level_colors.get(level, HexColor('#333333'))
            
            # 问题编号和标题
            item_title_style = ParagraphStyle(
                f'ItemTitle_{i}',
                parent=heading_style,
                fontSize=12,
                textColor=level_color,
                spaceBefore=4*mm,
                spaceAfter=2*mm
            )
            story.append(Paragraph(f'问题 {i} 【{level}风险】{item["category"]}', item_title_style))
            
            # 问题描述
            story.append(Paragraph(f'<b>问题描述：</b>{item["description"]}', label_style))
            
            # 原文引用
            if item.get('original_text'):
                story.append(Paragraph(f'<b>原文引用：</b>{item["original_text"]}', label_style))
            
            # 法规依据
            if item.get('legal_basis'):
                story.append(Paragraph(f'<b>法规依据：</b>{item["legal_basis"]}', label_style))
            
            # 修改建议
            if item.get('suggestion'):
                story.append(Paragraph(f'<b>修改建议：</b>{item["suggestion"]}', label_style))
            
            story.append(Spacer(1, 2*mm))
    
    # 免责声明
    story.append(Spacer(1, 5*mm))
    story.append(Paragraph('四、免责声明', heading_style))
    disclaimer_style = ParagraphStyle(
        'Disclaimer',
        parent=normal_style,
        fontSize=8,
        textColor=HexColor('#999999'),
        leading=12
    )
    story.append(Paragraph(
        '本报告由NEXUS招标文件合规检测系统自动生成，检测结果仅供参考。'
        '系统采用规则引擎与人工智能大语言模型相结合的方式进行分析，'
        '可能存在遗漏或误判。招标文件的最终合规性应以法律法规和行政监督部门的认定为准。'
        '建议在重要项目招标前，咨询专业法律顾问进行人工审查。',
        disclaimer_style
    ))
    
    # 构建PDF
    doc.build(story)


# ============================================================
# Flask路由
# ============================================================

@app.route('/')
def index():
    """首页 - 功能导航"""
    return render_template('index.html')


@app.route('/check')
def check_page():
    """合规检测页面"""
    return render_template('check.html')


@app.route('/api/check', methods=['POST'])
def check_document():
    """上传文件并返回检测结果"""
    # 检查文件
    if 'file' not in request.files:
        return jsonify({'error': '请上传文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': f'不支持的文件格式，请上传 {", ".join(ALLOWED_EXTENSIONS)} 格式'}), 400
    
    # 检查文件大小
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    if file_size > MAX_CONTENT_LENGTH:
        return jsonify({'error': '文件大小超过16MB限制'}), 400
    
    try:
        # 保存文件到临时目录
        filename = file.filename
        file_path = os.path.join(UPLOAD_FOLDER, f'tender_check_{uuid.uuid4().hex}_{filename}')
        file.save(file_path)
        
        print(f'[INFO] 开始处理文件: {filename}')
        
        # 1. 解析文档
        text = parse_document(file_path, filename)
        if not text or len(text.strip()) < 10:
            return jsonify({'error': '文件内容为空或无法解析'}), 400
        
        print(f'[INFO] 文档解析完成，文本长度: {len(text)} 字符')
        
        # 2. 规则引擎检测
        print('[INFO] 开始规则引擎检测...')
        rule_items = run_rule_check(text)
        print(f'[INFO] 规则引擎检测完成，发现 {len(rule_items)} 个问题')
        
        # 3. LLM深度分析
        llm_items = call_llm(text)
        
        # 4. 合并去重
        all_items = merge_and_deduplicate(rule_items, llm_items)
        
        # 5. 生成报告
        report_id = uuid.uuid4().hex
        summary = {
            'total': len(all_items),
            'high': sum(1 for item in all_items if item['level'] == '高'),
            'medium': sum(1 for item in all_items if item['level'] == '中'),
            'low': sum(1 for item in all_items if item['level'] == '低'),
        }
        
        report = {
            'id': report_id,
            'filename': filename,
            'check_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'summary': summary,
            'items': all_items,
            'text_length': len(text),
            'llm_enabled': bool(SILICONFLOW_API_KEY)
        }
        
        # 存储报告
        REPORTS[report_id] = report
        
        # 清理临时文件
        try:
            os.remove(file_path)
        except OSError:
            pass
        
        print(f'[INFO] 检测完成: 总计{summary["total"]}个问题 (高{summary["high"]}/中{summary["medium"]}/低{summary["low"]})')
        
        return jsonify(report)
        
    except Exception as e:
        traceback.print_exc()
        # 清理临时文件
        try:
            os.remove(file_path)
        except (OSError, UnboundLocalError):
            pass
        return jsonify({'error': f'处理文件时出错: {str(e)}'}), 500


@app.route('/api/report/<report_id>')
def get_report(report_id):
    """获取检测报告（支持JSON和PDF下载）"""
    if report_id not in REPORTS:
        return jsonify({'error': '报告不存在'}), 404
    
    report = REPORTS[report_id]
    
    # 如果请求参数中有 format=pdf，则返回PDF文件
    if request.args.get('format') == 'pdf':
        pdf_path = os.path.join(UPLOAD_FOLDER, f'report_{report_id}.pdf')
        try:
            generate_pdf_report(report, pdf_path)
            return send_file(
                pdf_path,
                as_attachment=True,
                download_name=f'合规检测报告_{report["filename"]}_{report["check_time"][:10]}.pdf',
                mimetype='application/pdf'
            )
        except Exception as e:
            traceback.print_exc()
            return jsonify({'error': f'生成PDF报告失败: {str(e)}'}), 500
    
    # 默认返回JSON
    return jsonify(report)


@app.route('/api/health')
def health_check():
    """健康检查接口"""
    return jsonify({
        'status': 'ok',
        'llm_enabled': bool(SILICONFLOW_API_KEY),
        'model': LLM_MODEL if SILICONFLOW_API_KEY else None
    })


# ============================================================
# 投标文件生成路由
# ============================================================

@app.route('/bid')
def bid_page():
    """投标文件生成页面"""
    return render_template('bid.html')


@app.route('/api/parse-tender', methods=['POST'])
def parse_tender():
    """上传招标文件，返回AI解析出的需求结构JSON"""
    if 'file' not in request.files:
        return jsonify({'error': '请上传招标文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': f'不支持的文件格式，请上传 {", ".join(ALLOWED_EXTENSIONS)} 格式'}), 400

    # 文件大小检查（招标文件限制16MB）
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    if file_size > MAX_CONTENT_LENGTH:
        return jsonify({'error': '招标文件大小超过16MB限制'}), 400

    try:
        # 保存文件
        filename = secure_filename(file.filename) or 'tender.txt'
        file_path = os.path.join(UPLOAD_FOLDER, f'tender_parse_{uuid.uuid4().hex}_{filename}')
        file.save(file_path)

        print(f'[INFO] 开始解析招标文件: {file.filename}')

        # 解析文档文本
        text = parse_document(file_path, file.filename)
        if not text or len(text.strip()) < 10:
            return jsonify({'error': '文件内容为空或无法解析'}), 400

        print(f'[INFO] 文档解析完成，文本长度: {len(text)} 字符')

        # 调用AI解析招标需求
        requirements = parse_tender_document(text)

        # 创建会话
        session_id = uuid.uuid4().hex
        BID_SESSIONS[session_id] = {
            'tender_text': text[:5000],  # 保留前5000字用于后续生成
            'requirements': requirements,
            'company_info': None,
            'materials': [],
            'created_at': datetime.now().isoformat(),
        }

        # 清理临时文件
        try:
            os.remove(file_path)
        except OSError:
            pass

        print(f'[INFO] 招标文件解析完成，会话ID: {session_id}')
        return jsonify({
            'session_id': session_id,
            'requirements': requirements,
        })

    except Exception as e:
        traceback.print_exc()
        try:
            os.remove(file_path)
        except (OSError, UnboundLocalError):
            pass
        return jsonify({'error': f'解析招标文件时出错: {str(e)}'}), 500


@app.route('/api/upload-material', methods=['POST'])
def upload_material():
    """上传佐证材料文件"""
    if 'file' not in request.files:
        return jsonify({'error': '请上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    material_name = request.form.get('material_name', '其他材料')
    session_id = request.form.get('session_id', 'default')

    filename = file.filename
    # 判断文件类型
    file_type = get_file_type(filename)
    if not file_type:
        return jsonify({'error': f'不支持的文件格式: {filename}'}), 400

    # 文件大小检查
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    size_limit = get_file_size_limit(file_type)
    if file_size > size_limit:
        limit_mb = size_limit / (1024 * 1024)
        return jsonify({'error': f'文件大小超过限制（{FILE_TYPE_LABELS[file_type]}限{limit_mb:.0f}MB）'}), 400

    try:
        # 安全文件名
        safe_name = secure_filename(filename) or f'material_{uuid.uuid4().hex}'
        ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        # 保留中文文件名（secure_filename会移除中文）
        if not safe_name or safe_name == '':
            safe_name = f'material_{uuid.uuid4().hex}.{ext}'

        # 存储路径：upload/{session_id}/{文件类型中文}/
        type_label = FILE_TYPE_LABELS.get(file_type, 'other')
        save_dir = os.path.join(UPLOAD_DIR, session_id, type_label)
        os.makedirs(save_dir, exist_ok=True)
        # 避免文件名冲突
        unique_name = f'{uuid.uuid4().hex[:8]}_{safe_name}'
        save_path = os.path.join(save_dir, unique_name)
        file.save(save_path)

        # 生成预览URL（仅图片）
        preview_url = None
        if file_type == 'image':
            preview_url = f'/api/material-preview/{session_id}/{type_label}/{unique_name}'

        # 记录到会话
        if session_id in BID_SESSIONS:
            BID_SESSIONS[session_id]['materials'].append({
                'filename': filename,
                'saved_name': unique_name,
                'material_name': material_name,
                'file_type': file_type,
                'size': file_size,
                'path': save_path,
                'preview_url': preview_url,
            })

        print(f'[INFO] 材料上传: {filename} -> {material_name} [{file_type}, {file_size}B]')

        return jsonify({
            'filename': filename,
            'material_name': material_name,
            'category': material_name,
            'file_type': file_type,
            'size': file_size,
            'path': save_path,
            'preview_url': preview_url,
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'上传文件时出错: {str(e)}'}), 500


@app.route('/api/material-preview/<session_id>/<type_label>/<filename>')
def material_preview(session_id, type_label, filename):
    """预览佐证材料（图片）"""
    file_path = os.path.join(UPLOAD_DIR, session_id, type_label, filename)
    if not os.path.exists(file_path):
        return jsonify({'error': '文件不存在'}), 404
    return send_file(file_path)


@app.route('/api/generate-bid', methods=['POST'])
def generate_bid():
    """提交企业信息+招标需求，生成投标文件.docx"""
    data = request.get_json()
    if not data:
        return jsonify({'error': '请提交JSON数据'}), 400

    session_id = data.get('session_id', '')
    company_info = data.get('company_info', {})
    uploaded_materials = data.get('uploaded_materials', [])

    if not company_info.get('company_name'):
        return jsonify({'error': '请填写企业名称'}), 400

    # 获取招标需求
    tender_req = None
    if session_id and session_id in BID_SESSIONS:
        tender_req = BID_SESSIONS[session_id].get('requirements')
        # 更新材料信息
        BID_SESSIONS[session_id]['company_info'] = company_info

    if not tender_req:
        return jsonify({'error': '招标需求不存在，请重新上传招标文件'}), 400

    try:
        print(f'[INFO] 开始生成投标文件: {company_info["company_name"]}')

        # 1. 调用AI生成投标内容
        bid_content = generate_bid_content(tender_req, company_info)

        # 2. 计算材料匹配状态
        material_status = get_material_status(tender_req, uploaded_materials)

        # 3. 生成docx文件
        company_short = company_info['company_name'][:20]
        project_short = tender_req.get('project_name', '项目')[:20]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f'投标文件_{company_short}_{project_short}_{timestamp}.docx'
        # 清理文件名中的非法字符
        output_filename = ''.join(c for c in output_filename if c not in r'\/:*?"<>|')
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        generate_bid_docx(tender_req, company_info, bid_content, material_status, output_path)

        # 4. 获取缺失材料
        missing = get_missing_materials(material_status)

        # 5. 统计信息
        provided_count = sum(1 for m in material_status if m['status'] == 'provided')
        tech_count = len(bid_content.get('technical_responses', []))

        print(f'[INFO] 投标文件生成完成: {output_filename}')

        return jsonify({
            'filename': output_filename,
            'missing_materials': missing,
            'summary': {
                'technical_count': tech_count,
                'provided_count': provided_count,
                'missing_count': len(missing),
                'sections': 10,
                'llm_disabled': bid_content.get('_llm_disabled', False),
            }
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'生成投标文件时出错: {str(e)}'}), 500


@app.route('/api/download-bid/<path:filename>')
def download_bid(filename):
    """下载生成的投标文件"""
    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        return jsonify({'error': '文件不存在，可能已过期'}), 404

    # 中文文件名需要处理
    download_name = os.path.basename(filename)
    return send_file(
        file_path,
        as_attachment=True,
        download_name=download_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    # 启动提示
    print('=' * 60)
    print('  NEXUS · 招标文件合规检测 + AI投标文件生成')
    print('=' * 60)
    print(f'  LLM深度分析: {"已启用" if SILICONFLOW_API_KEY else "未启用（仅规则引擎/模板模式）"}')
    if SILICONFLOW_API_KEY:
        print(f'  模型: {LLM_MODEL}')
    else:
        print('  提示: 设置 SILICONFLOW_API_KEY 环境变量以启用AI功能')
    print(f'  合规检测: http://127.0.0.1:5000/')
    print(f'  投标生成: http://127.0.0.1:5000/bid')
    print(f'  健康检查: http://127.0.0.1:5000/api/health')
    print('=' * 60)

    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
