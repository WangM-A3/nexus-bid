# -*- coding: utf-8 -*-
"""
投标文件生成后端 —— 招标解析 + AI生成 + 材料管理 + docx输出
功能：解析招标文件提取需求、管理佐证材料、调用LLM生成投标内容、输出.docx文件
"""

import os
import json
import uuid
import traceback
from datetime import datetime

import requests
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from bid_templates import (
    MATERIAL_CATEGORIES,
    build_tender_parse_prompt,
    build_bid_generation_prompt,
    get_bid_letter_template,
    get_commitment_templates,
)

# 硅基流动API配置（延迟读取，由app.py通过os.environ.setdefault设置）
SILICONFLOW_API_URL = 'https://api.siliconflow.cn/v1/chat/completions'
LLM_MODEL = 'Qwen/Qwen2.5-72B-Instruct'

# 存储目录
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, 'upload')
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')

# 内存存储会话数据（无需数据库）
BID_SESSIONS = {}


# ============================================================
# LLM调用
# ============================================================

def call_llm(system_prompt, user_prompt, max_tokens=4096):
    """
    调用硅基流动API
    参数: system_prompt, user_prompt
    返回: LLM返回的文本内容，失败返回None
    """
    # 延迟读取API Key（由app.py通过os.environ.setdefault设置）
    api_key = os.environ.get('SILICONFLOW_API_KEY', '')
    if not api_key:
        print('[WARN] 未配置SILICONFLOW_API_KEY，无法调用LLM')
        return None

    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}'
    }

    payload = {
        'model': LLM_MODEL,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.3,
        'max_tokens': max_tokens,
        'response_format': {'type': 'json_object'}
    }

    try:
        print('[INFO] 调用硅基流动API...')
        response = requests.post(SILICONFLOW_API_URL, headers=headers, json=payload, timeout=180)
        response.raise_for_status()

        result = response.json()
        content = result['choices'][0]['message']['content']
        return content

    except requests.exceptions.Timeout:
        print('[WARN] LLM调用超时')
        return None
    except Exception as e:
        print(f'[WARN] LLM调用失败: {e}')
        return None


# ============================================================
# 招标文件解析
# ============================================================

def parse_tender_document(text):
    """
    解析招标文件，提取结构化需求
    参数: text - 招标文件全文
    返回: 解析结果dict，包含项目信息、要求、材料清单
    """
    system_prompt, user_prompt = build_tender_parse_prompt(text)
    content = call_llm(system_prompt, user_prompt, max_tokens=4096)

    if not content:
        # LLM不可用时的降级处理：返回基础结构
        print('[INFO] LLM不可用，返回基础解析结果')
        return _fallback_parse(text)

    try:
        result = json.loads(content)
        # 确保字段完整
        result.setdefault('project_name', '')
        result.setdefault('project_budget', '')
        result.setdefault('deadline', '')
        result.setdefault('qualifications', [])
        result.setdefault('technical_requirements', [])
        result.setdefault('commercial_requirements', [])
        result.setdefault('evaluation_criteria', {})
        result.setdefault('key_dates', {})
        result.setdefault('required_materials', [])

        # 如果LLM没有提取到材料清单，用默认分类补充
        if not result['required_materials']:
            result['required_materials'] = _get_default_materials()

        print(f'[INFO] 招标文件解析完成: {result["project_name"]}')
        return result

    except json.JSONDecodeError as e:
        print(f'[WARN] LLM返回JSON解析失败: {e}')
        return _fallback_parse(text)


def _fallback_parse(text):
    """LLM不可用时的降级解析：简单关键词提取 + 默认材料清单"""
    project_name = ''
    for line in text.split('\n')[:20]:
        line = line.strip()
        if '项目名称' in line or '项目' in line:
            project_name = line
            break

    return {
        'project_name': project_name or '（请手动填写项目名称）',
        'project_budget': '',
        'deadline': '',
        'qualifications': [],
        'technical_requirements': ['（LLM不可用，请手动填写技术要求）'],
        'commercial_requirements': [],
        'evaluation_criteria': {},
        'key_dates': {},
        'required_materials': _get_default_materials(),
        '_llm_disabled': True
    }


def _get_default_materials():
    """获取默认佐证材料清单（当LLM未提取时使用）"""
    defaults = ['business_license', 'legal_rep_id', 'auth_letter',
                'performance_contract', 'financial_report',
                'social_security', 'team_cert']
    return [
        {
            'name': MATERIAL_CATEGORIES[key]['name'],
            'type': MATERIAL_CATEGORIES[key]['type'],
            'desc': MATERIAL_CATEGORIES[key]['desc']
        }
        for key in defaults
    ]


# ============================================================
# 投标文件内容生成
# ============================================================

def generate_bid_content(tender_req, company_info):
    """
    调用LLM生成投标文件核心内容
    参数: tender_req - 招标需求, company_info - 企业信息
    返回: 生成的内容dict
    """
    system_prompt, user_prompt = build_bid_generation_prompt(tender_req, company_info)
    content = call_llm(system_prompt, user_prompt, max_tokens=4096)

    if not content:
        # 降级：使用模板生成
        print('[INFO] LLM不可用，使用模板生成')
        return _fallback_generate(tender_req, company_info)

    try:
        result = json.loads(content)
        result.setdefault('bid_letter', '')
        result.setdefault('technical_responses', [])
        result.setdefault('commercial_response', '')
        print(f'[INFO] 投标内容生成完成: 技术响应{len(result["technical_responses"])}条')
        return result

    except json.JSONDecodeError as e:
        print(f'[WARN] LLM返回JSON解析失败: {e}')
        return _fallback_generate(tender_req, company_info)


def _fallback_generate(tender_req, company_info):
    """LLM不可用时的降级生成"""
    company_name = company_info.get('company_name', '________')
    project_name = tender_req.get('project_name', '________')
    bid_price = company_info.get('bid_price', '')
    delivery = company_info.get('delivery_period', '')

    # 用模板生成投标函
    bid_letter = get_bid_letter_template(company_name, project_name, bid_price, delivery)

    # 技术响应：逐条简单响应
    tech_reqs = tender_req.get('technical_requirements', [])
    tech_responses = []
    for req in tech_reqs:
        tech_responses.append({
            'requirement': req,
            'response': f'我方完全响应并满足该项要求。具体实施方案将根据项目实际情况制定详细计划，确保达到或优于招标文件要求的标准。'
        })

    # 商务响应
    commercial = '我方将严格按照招标文件的商务要求执行，具体报价详见投标一览表。付款方式、售后服务等均按照招标文件要求响应。'

    return {
        'bid_letter': bid_letter,
        'technical_responses': tech_responses,
        'commercial_response': commercial,
        '_llm_disabled': True
    }


# ============================================================
# 佐证材料管理
# ============================================================

def get_material_status(tender_req, uploaded_materials):
    """
    计算佐证材料的匹配状态
    参数: tender_req - 招标需求(含required_materials), uploaded_materials - 已上传材料列表
    返回: 材料状态列表 [{name, type, desc, status: provided/missing, files: [...]}]
    """
    required = tender_req.get('required_materials', [])
    status_list = []

    for req_mat in required:
        mat_name = req_mat.get('name', '')
        mat_type = req_mat.get('type', 'must')
        mat_desc = req_mat.get('desc', '')

        # 查找已上传的匹配材料（按名称模糊匹配）
        matched_files = []
        for up_mat in uploaded_materials:
            up_category = up_mat.get('category', '')
            up_filename = up_mat.get('filename', '')
            # 匹配规则：材料分类key匹配 或 文件名包含材料名称
            if _match_material(mat_name, up_category, up_filename):
                matched_files.append(up_mat)

        status_list.append({
            'name': mat_name,
            'type': mat_type,
            'desc': mat_desc,
            'status': 'provided' if matched_files else 'missing',
            'files': matched_files
        })

    return status_list


def _match_material(required_name, uploaded_category, uploaded_filename):
    """判断已上传材料是否匹配需求材料"""
    # 通过材料分类key匹配
    for key, cat in MATERIAL_CATEGORIES.items():
        if cat['name'] == required_name and key == uploaded_category:
            return True

    # 通过名称模糊匹配
    if required_name and required_name in uploaded_filename:
        return True

    return False


def get_missing_materials(material_status):
    """获取缺失的材料列表"""
    return [m for m in material_status if m['status'] == 'missing']


# ============================================================
# docx投标文件生成
# ============================================================

def _set_cell_font(cell, font_name='宋体', font_size=10.5, bold=False):
    """设置表格单元格字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _add_heading(doc, text, level=1):
    """添加标题（设置中文字体）"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def _add_paragraph(doc, text, font_name='宋体', font_size=12, bold=False, alignment=None):
    """添加段落（设置中文字体）"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def generate_bid_docx(tender_req, company_info, bid_content, material_status, output_path):
    """
    生成投标文件.docx
    参数:
        tender_req - 招标需求
        company_info - 企业信息
        bid_content - LLM生成的投标内容
        material_status - 材料匹配状态
        output_path - 输出文件路径
    返回: output_path
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    company_name = company_info.get('company_name', '________')
    project_name = tender_req.get('project_name', '________')
    bid_price = company_info.get('bid_price', '（详见投标一览表）')
    delivery = company_info.get('delivery_period', '（详见投标一览表）')
    today = datetime.now().strftime('%Y年%m月%d日')

    # ====== 1. 封面 ======
    for _ in range(6):
        doc.add_paragraph()
    _add_paragraph(doc, project_name, font_name='黑体', font_size=26, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _add_paragraph(doc, '投 标 文 件', font_name='黑体', font_size=36, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        doc.add_paragraph()
    _add_paragraph(doc, f'投标人：{company_name}', font_name='宋体', font_size=16,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f'日期：{today}', font_name='宋体', font_size=16,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ====== 2. 投标函 ======
    _add_heading(doc, '一、投标函', level=1)
    bid_letter = bid_content.get('bid_letter', '')
    if bid_letter:
        for line in bid_letter.split('\n'):
            if line.strip():
                _add_paragraph(doc, line.strip(), font_size=12)
    else:
        letter = get_bid_letter_template(company_name, project_name, bid_price, delivery)
        for line in letter.split('\n'):
            if line.strip():
                _add_paragraph(doc, line.strip(), font_size=12)
    doc.add_page_break()

    # ====== 3. 投标一览表 ======
    _add_heading(doc, '二、投标一览表', level=1)
    table = doc.add_table(rows=6, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows_data = [
        ('项目名称', project_name),
        ('投标人', company_name),
        ('投标报价（元）', bid_price or '（后续填写）'),
        ('工期/交货期', delivery or '（后续填写）'),
        ('质量保证', '符合招标文件要求，达到合格标准'),
        ('投标有效期', '自投标截止日起90个日历天'),
    ]
    for i, (label, value) in enumerate(rows_data):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = label
        cell_value.text = str(value)
        _set_cell_font(cell_label, bold=True)
        _set_cell_font(cell_value)
        # 设置列宽
        cell_label.width = Cm(5)
        cell_value.width = Cm(11)

    doc.add_page_break()

    # ====== 4. 资格证明文件 ======
    _add_heading(doc, '三、资格证明文件', level=1)
    _add_paragraph(doc, '投标人提交以下资格证明文件，证明其具备承担本项目的能力和条件：', font_size=12)

    # 资质证书清单
    qualifications = company_info.get('qualifications', [])
    if qualifications:
        _add_paragraph(doc, '（一）资质证书清单', font_name='黑体', font_size=14, bold=True)
        qual_table = doc.add_table(rows=len(qualifications) + 1, cols=3, style='Table Grid')
        qual_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 表头
        headers = ['序号', '资质证书名称', '备注']
        for j, h in enumerate(headers):
            cell = qual_table.cell(0, j)
            cell.text = h
            _set_cell_font(cell, bold=True)
        for i, qual in enumerate(qualifications):
            qual_table.cell(i + 1, 0).text = str(i + 1)
            qual_table.cell(i + 1, 1).text = qual
            qual_table.cell(i + 1, 2).text = '有效'
            for j in range(3):
                _set_cell_font(qual_table.cell(i + 1, j))

    _add_paragraph(doc, '')
    _add_paragraph(doc, '（二）相关证明文件', font_name='黑体', font_size=14, bold=True)
    _add_paragraph(doc, '1. 营业执照副本复印件（详见附件）', font_size=12)
    _add_paragraph(doc, '2. 法定代表人身份证复印件（详见附件）', font_size=12)
    _add_paragraph(doc, '3. 授权委托书（详见附件）', font_size=12)
    if qualifications:
        _add_paragraph(doc, '4. 资质证书复印件（详见附件）', font_size=12)

    doc.add_page_break()

    # ====== 5. 技术方案 ======
    _add_heading(doc, '四、技术方案', level=1)
    _add_paragraph(doc, '我方已仔细阅读招标文件中的技术要求，现逐条响应如下：', font_size=12)

    tech_responses = bid_content.get('technical_responses', [])
    tech_reqs = tender_req.get('technical_requirements', [])

    if tech_responses:
        for i, resp in enumerate(tech_responses, 1):
            req_text = resp.get('requirement', '')
            resp_text = resp.get('response', '')
            _add_heading(doc, f'4.{i} 技术要求：{req_text[:50]}{"..." if len(req_text) > 50 else ""}', level=2)
            _add_paragraph(doc, f'【招标要求】{req_text}', font_size=11)
            _add_paragraph(doc, f'【我方响应】{resp_text}', font_size=12)
            doc.add_paragraph()
    else:
        # 无LLM响应时逐条列出要求
        for i, req in enumerate(tech_reqs, 1):
            _add_heading(doc, f'4.{i} {req[:50]}{"..." if len(req) > 50 else ""}', level=2)
            _add_paragraph(doc, f'【招标要求】{req}', font_size=11)
            _add_paragraph(doc, '【我方响应】我方完全响应并满足该项技术要求。', font_size=12)
            doc.add_paragraph()

    doc.add_page_break()

    # ====== 6. 商务方案 ======
    _add_heading(doc, '五、商务方案', level=1)

    commercial = bid_content.get('commercial_response', '')
    if commercial:
        for line in commercial.split('\n'):
            if line.strip():
                _add_paragraph(doc, line.strip(), font_size=12)
    else:
        _add_paragraph(doc, '我方将严格按照招标文件的商务要求执行，具体内容如下：', font_size=12)

    doc.add_paragraph()
    # 商务要求响应表
    comm_reqs = tender_req.get('commercial_requirements', [])
    if comm_reqs:
        _add_paragraph(doc, '商务要求逐条响应：', font_name='黑体', font_size=14, bold=True)
        comm_table = doc.add_table(rows=len(comm_reqs) + 1, cols=3, style='Table Grid')
        comm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['序号', '招标商务要求', '我方响应']
        for j, h in enumerate(headers):
            cell = comm_table.cell(0, j)
            cell.text = h
            _set_cell_font(cell, bold=True)
        for i, req in enumerate(comm_reqs):
            comm_table.cell(i + 1, 0).text = str(i + 1)
            comm_table.cell(i + 1, 1).text = req
            comm_table.cell(i + 1, 2).text = '完全响应'
            for j in range(3):
                _set_cell_font(comm_table.cell(i + 1, j))

    doc.add_page_break()

    # ====== 7. 项目团队 ======
    _add_heading(doc, '六、项目团队', level=1)
    team_members = company_info.get('team_members', [])
    if team_members:
        _add_paragraph(doc, '我方为本项目组建专业项目团队，人员配置如下：', font_size=12)
        team_table = doc.add_table(rows=len(team_members) + 1, cols=2, style='Table Grid')
        team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        team_table.cell(0, 0).text = '序号'
        team_table.cell(0, 1).text = '人员配置'
        _set_cell_font(team_table.cell(0, 0), bold=True)
        _set_cell_font(team_table.cell(0, 1), bold=True)
        for i, member in enumerate(team_members):
            team_table.cell(i + 1, 0).text = str(i + 1)
            team_table.cell(i + 1, 1).text = member
            _set_cell_font(team_table.cell(i + 1, 0))
            _set_cell_font(team_table.cell(i + 1, 1))
    else:
        _add_paragraph(doc, '项目团队人员配置详见附件。', font_size=12)

    doc.add_page_break()

    # ====== 8. 业绩证明 ======
    _add_heading(doc, '七、业绩证明', level=1)
    performances = company_info.get('performances', [])
    if performances:
        _add_paragraph(doc, '我方近年完成的同类项目业绩如下：', font_size=12)
        perf_table = doc.add_table(rows=len(performances) + 1, cols=2, style='Table Grid')
        perf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        perf_table.cell(0, 0).text = '序号'
        perf_table.cell(0, 1).text = '项目业绩'
        _set_cell_font(perf_table.cell(0, 0), bold=True)
        _set_cell_font(perf_table.cell(0, 1), bold=True)
        for i, perf in enumerate(performances):
            perf_table.cell(i + 1, 0).text = str(i + 1)
            perf_table.cell(i + 1, 1).text = perf
            _set_cell_font(perf_table.cell(i + 1, 0))
            _set_cell_font(perf_table.cell(i + 1, 1))
        _add_paragraph(doc, '')
        _add_paragraph(doc, '注：以上业绩证明材料详见附件。', font_size=11)
    else:
        _add_paragraph(doc, '业绩证明材料详见附件。', font_size=12)

    doc.add_page_break()

    # ====== 9. 承诺函 ======
    _add_heading(doc, '八、承诺函', level=1)
    commitments = get_commitment_templates(company_name)
    for commit in commitments:
        _add_heading(doc, commit['title'], level=2)
        content = commit['content'].replace('________', project_name)
        for line in content.split('\n'):
            if line.strip():
                _add_paragraph(doc, line.strip(), font_size=12)
        doc.add_paragraph()

    doc.add_page_break()

    # ====== 10. 附件清单 ======
    _add_heading(doc, '九、附件清单', level=1)
    _add_paragraph(doc, '本投标文件包含以下附件：', font_size=12)

    # 收集所有已提供材料作为附件
    attachment_idx = 1
    for mat in material_status:
        mat_name = mat['name']
        files = mat.get('files', [])
        if files:
            for f in files:
                ftype = f.get('file_type', 'document')
                ftype_label = {'document': '文档', 'image': '图片', 'video': '视频', 'archive': '压缩包'}.get(ftype, '文件')
                ext = f.get('filename', '').rsplit('.', 1)[-1] if '.' in f.get('filename', '') else ''
                ext_info = f'（{ext.upper()}格式）' if ext else ''
                _add_paragraph(doc, f'附件{attachment_idx}：{mat_name}{ext_info}', font_size=12)
                attachment_idx += 1
        else:
            # 缺失材料也列出，标注待补充
            _add_paragraph(doc, f'附件{attachment_idx}：{mat_name}（待补充）', font_size=12, bold=False)
            attachment_idx += 1

    # 保存文件
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'[INFO] 投标文件已生成: {output_path}')
    return output_path
